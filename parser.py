import os
import pytesseract
import pdfplumber
#For windows install Poppler
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma

# For Windows, set the path to Tesseract and Poppler, needs to be installed
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH=r'C:\Program Files\poppler-windows-24.08.0-0'
FOLDER_PATH = "./PoliciesForTheTask"
load_dotenv()

class Parser:

    def __init__(self):
        self.directory = None

    def parse_pdf(self, file_path):
        """Extract text from a PDF (either searchable or scanned)."""
        text = ""

        try:
            # Try extracting text from searchable PDFs
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""  # Ensure it doesn't break on None

            if text.strip():  # If we got text, return it
                return text

            # If not, assume it's a scanned PDF and use OCR
            images = convert_from_path(file_path, poppler_path=POPPLER_PATH)
            text = "\n".join([pytesseract.image_to_string(img) for img in images])
            return text

        except Exception as e:
            print(f"Error parsing PDF {file_path}: {e}")
            return None

    def parse_image(self, file_path):
        """Extract text from an image using OCR."""
        try:
            image = Image.open(file_path)
            return pytesseract.image_to_string(image)
        except Exception as e:
            print(f"Error parsing image {file_path}: {e}")
            return None

    def parse_docx(self, file_path):
        """Extract text from a Word document."""
        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Error parsing Word document {file_path}: {e}")
            return None

    def parse_file(self,file_path):
        """Determine file type and parse accordingly."""
        ext = file_path.lower().split('.')[-1]

        if ext == "pdf":
            return self.parse_pdf(file_path)
        elif ext in ["png", "jpg", "jpeg", "tiff"]:
            return self.parse_image(file_path)
        elif ext in ["docx"]:
            return self.parse_docx(file_path)
        else:
            print(f"Unsupported file format: {file_path}")
            return None

    def process_directory(self, directory):
        """Parse all supported files in a directory and return LangChain-compatible Documents."""
        langchain_docs = []

        for root, _, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                text = self.parse_file(file_path)

                if text:
                    text = self.clean_text(text)
                    langchain_docs.append(Document(page_content=text, metadata={"source": file_path}))

        return langchain_docs

    def clean_text(self, text: str) -> str:

        # Remove extra newlines, headers, bullet points etc.
        text = re.sub(r"\n+", " ", text)
        text = re.sub(r"\s{2,}", " ", text)
        text = re.sub(r"[-●•◦○]", "", text)
        text = re.sub(r"Page \d+|Table of Contents|\.{3,}", "", text)
        return text.strip()

    def store_in_chroma(self, directory, embeddings):
        self.docs = self.process_directory(directory)
        print(f"Loaded {len(self.docs)} documents into LangChain.")

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # approx. 1000 characters per chunk (adjust as needed)
            chunk_overlap=200
        )
        split_docs = text_splitter.split_documents(self.docs)
        print(f"Split into {len(split_docs)} chunks.")

        persist_directory = "chroma_db"  # directory for persistent storage

        # Create (or load) the Chroma vector store.
        vectorstore = Chroma.from_documents(split_docs, embeddings, persist_directory=persist_directory)
        print("Documents stored in Chroma.")

        return vectorstore

if __name__ == "__main__":
    ps = Parser()
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    ps.store_in_chroma(FOLDER_PATH, embeddings)
